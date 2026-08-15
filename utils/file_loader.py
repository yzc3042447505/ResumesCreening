"""
文件读取工具模块
功能：将PDF、DOCX等格式的简历文件转换为纯文本
说明：支持多种常见简历格式，统一输出纯文本供后续Agent解析
"""

import os
from typing import Tuple


def extract_text_from_pdf(file_path: str) -> str:
    """
    从PDF文件中提取纯文本

    参数:
        file_path: PDF文件的完整路径

    返回:
        str: 提取出的纯文本内容

    注意:
        只能提取文本型PDF，扫描件/图片型PDF无法提取（需要OCR，MVP阶段不支持）
    """
    try:
        from pypdf import PdfReader

        # 创建PDF读取器
        reader = PdfReader(file_path)
        text_list = []

        # 遍历每一页，提取文本
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_list.append(page_text)

        # 把所有页的文本拼接起来，用换行分隔
        full_text = "\n".join(text_list)
        return full_text.strip()

    except ImportError:
        raise Exception("缺少pypdf库，请运行 pip install pypdf")
    except Exception as e:
        raise Exception(f"PDF文件读取失败: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    从DOCX文件中提取纯文本

    参数:
        file_path: DOCX文件的完整路径

    返回:
        str: 提取出的纯文本内容
    """
    try:
        from docx import Document

        # 打开Word文档
        doc = Document(file_path)
        text_list = []

        # 读取所有段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_list.append(para.text)

        # 读取所有表格中的内容（有些简历会用表格排版）
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_list.append(" | ".join(row_texts))

        # 拼接所有文本
        full_text = "\n".join(text_list)
        return full_text.strip()

    except ImportError:
        raise Exception("缺少python-docx库，请运行 pip install python-docx")
    except Exception as e:
        raise Exception(f"DOCX文件读取失败: {str(e)}")


def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    通用文件读取接口，自动识别文件类型并提取文本

    参数:
        file_path: 文件的完整路径

    返回:
        tuple: (文件类型, 提取的文本内容)
        文件类型: 'pdf' / 'docx' / 'unknown'
    """
    # 获取文件扩展名，转小写
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        return "pdf", text

    elif file_ext == ".docx":
        text = extract_text_from_docx(file_path)
        return "docx", text

    elif file_ext == ".doc":
        # .doc 是旧版Word格式，python-docx不支持
        # MVP阶段先不支持，提示用户转成docx
        raise Exception("暂不支持.doc格式，请先另存为.docx格式后再上传")

    elif file_ext == ".txt":
        # 纯文本文件直接读取
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return "txt", text.strip()
        except UnicodeDecodeError:
            # 尝试用gbk编码读取（有些Windows下的txt文件是gbk编码）
            with open(file_path, "r", encoding="gbk") as f:
                text = f.read()
            return "txt", text.strip()

    else:
        raise Exception(f"不支持的文件格式: {file_ext}，目前支持 PDF、DOCX、TXT 格式")


def get_file_name(file_path: str) -> str:
    """
    从文件路径中提取文件名（不含扩展名）
    用于显示候选人姓名等（简历文件名通常是姓名）

    参数:
        file_path: 文件的完整路径

    返回:
        str: 文件名（不含扩展名）
    """
    base_name = os.path.basename(file_path)
    name_without_ext = os.path.splitext(base_name)[0]
    return name_without_ext


# 测试代码（直接运行本文件时执行）
